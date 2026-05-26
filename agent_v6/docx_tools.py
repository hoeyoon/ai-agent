from pathlib import Path
import re
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

from agent_v6.context import NS, WORD_NS, normalize_text, text_of
from agent_v6.xml_protocol import Action


W = f"{{{WORD_NS}}}"


class ToolError(RuntimeError):
    pass


def apply_actions(input_path: Path, actions: list[Action], output_path: Path) -> int:
    document_xml = read_document_xml(input_path)
    root = etree.fromstring(document_xml)
    changed = apply_actions_to_document_xml(root, actions)
    if changed == 0:
        raise ToolError("적용된 DOCX 수정이 없습니다.")
    remove_repeating_table_headers(root)
    write_docx_with_document_xml(input_path, output_path, root)
    return changed


def apply_actions_to_document_xml(root: etree._Element, actions: list[Action]) -> int:
    changed = 0
    for action in actions:
        if action.tool == "fill_template":
            for target, value in (action.fields or {}).items():
                changed += replace_table_value(root, target, value)
        elif action.tool == "replace_table_value":
            changed += replace_table_value(root, action.label, action.value)
        elif action.tool == "replace_paragraph":
            if action.index is None:
                raise ToolError("replace_paragraph index가 없습니다.")
            changed += replace_paragraph(root, action.index, action.value)
        elif action.tool == "replace_text":
            changed += replace_text(root, action.old, action.new)
        else:
            raise ToolError(f"지원하지 않는 tool입니다: {action.tool}")
    return changed


def read_document_xml(path: Path) -> bytes:
    with ZipFile(path) as archive:
        return archive.read("word/document.xml")


def write_docx_with_document_xml(input_path: Path, output_path: Path, root: etree._Element) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    new_document_xml = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    with ZipFile(input_path, "r") as source, ZipFile(output_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            if item.filename == "word/document.xml":
                target.writestr(item, new_document_xml)
            else:
                target.writestr(item, source.read(item.filename))


def replace_table_value(root: etree._Element, label: str, value: str) -> int:
    coordinate_change = replace_table_coordinate(root, label, value)
    if coordinate_change:
        return coordinate_change
    return replace_table_label(root, label, value)


def replace_table_coordinate(root: etree._Element, label: str, value: str) -> int:
    match = re.fullmatch(
        r"(?:table\[(?P<table>\d+)\]\s*)?row\[(?P<row>\d+)\]\s*cell\[(?P<cell>\d+)\]",
        normalize_text(label),
        flags=re.IGNORECASE,
    )
    if not match:
        return 0
    table = get_indexed(root.xpath("./w:body/w:tbl", namespaces=NS), int(match.group("table") or 0))
    if table is None:
        return 0
    row = get_indexed(table.findall("./w:tr", NS), int(match.group("row")))
    if row is None:
        return 0
    cell = get_indexed(row.findall("./w:tc", NS), int(match.group("cell")))
    if cell is None:
        return 0
    set_container_text(cell, value)
    return 1


def replace_table_label(root: etree._Element, label: str, value: str) -> int:
    changed = 0
    for table in root.xpath("./w:body/w:tbl", namespaces=NS):
        for row in table.findall("./w:tr", NS):
            cells = row.findall("./w:tc", NS)
            for index, cell in enumerate(cells[:-1]):
                if label_matches(text_of(cell), label):
                    set_container_text(cells[index + 1], value)
                    changed += 1
    return changed


def replace_paragraph(root: etree._Element, index: int, value: str) -> int:
    visible_index = 0
    for paragraph in root.xpath("./w:body/w:p", namespaces=NS):
        if not normalize_text(text_of(paragraph)):
            continue
        if visible_index == index:
            set_container_text(paragraph, value)
            return 1
        visible_index += 1
    return 0


def replace_text(root: etree._Element, old: str, new: str) -> int:
    if not old:
        return 0
    changed = 0
    for text_node in root.xpath(".//w:t", namespaces=NS):
        if text_node.text and old in text_node.text:
            text_node.text = text_node.text.replace(old, new)
            changed += 1
    return changed


def remove_repeating_table_headers(root: etree._Element) -> int:
    removed = 0
    for tbl_header in root.xpath(".//w:tblHeader", namespaces=NS):
        parent = tbl_header.getparent()
        if parent is not None:
            parent.remove(tbl_header)
            removed += 1
    return removed


def set_container_text(container: etree._Element, value: str) -> None:
    text_nodes = container.xpath(".//w:t", namespaces=NS)
    if text_nodes:
        text_nodes[0].text = value
        preserve_space(text_nodes[0])
        for text_node in text_nodes[1:]:
            text_node.text = ""
            preserve_space(text_node)
        return

    paragraph = first_or_create(container, "p")
    run = first_or_create(paragraph, "r")
    text_node = etree.SubElement(run, W + "t")
    text_node.text = value
    preserve_space(text_node)


def first_or_create(parent: etree._Element, local_name: str) -> etree._Element:
    child = parent.find(f"./w:{local_name}", NS)
    if child is not None:
        return child
    return etree.SubElement(parent, W + local_name)


def preserve_space(text_node: etree._Element) -> None:
    if text_node.text and (text_node.text[:1].isspace() or text_node.text[-1:].isspace()):
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def label_matches(cell_text: str, label: str) -> bool:
    return compact_key(cell_text) == compact_key(label)


def compact_key(value: str) -> str:
    return "".join(normalize_text(value).split()).lower()


def get_indexed(items: list[etree._Element], index: int) -> etree._Element | None:
    if index < 0 or index >= len(items):
        return None
    return items[index]


def write_document_xml(root: etree._Element, output_path: Path) -> None:
    document = Document()
    title = xml_text_of(root.find("title"))
    description = xml_text_of(root.find("description"))
    if title:
        document.add_heading(title, level=1)
    if description:
        document.add_paragraph(description)
    for section in root.findall("section"):
        heading = xml_text_of(section.find("heading"))
        if heading:
            document.add_heading(heading, level=2)
        for child in section:
            if child.tag == "paragraph" and xml_text_of(child):
                document.add_paragraph(xml_text_of(child))
    document.save(str(output_path))


def xml_text_of(element: etree._Element | None) -> str:
    if element is None:
        return ""
    return " ".join("".join(element.itertext()).split())

